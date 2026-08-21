"""分析结果导出模块。

支持将各类分析结果导出为：
- TXT（纯文本，便于阅读）
- CSV（表格数据，便于导入 Excel）
- JSON（结构化数据，便于二次开发）
- DOCX（Word 报告，便于撰写研究报告）

所有导出路径均通过 Tkinter 文件对话框获取。
"""

from __future__ import annotations

import csv
import json
import os
from datetime import datetime
from tkinter import filedialog, messagebox
from typing import Optional

from core.analyzer import BasicResult, SyntaxResult
from core.comparison import EnReadability, ZhReadability, Alignment
from core.linguistic_fingerprint import FingerprintResult


# --------------------------------------------------------------------------- #
# 通用对话框
# --------------------------------------------------------------------------- #


def _ask_path(title: str, default_ext: str, filetypes: list) -> Optional[str]:
    path = filedialog.asksaveasfilename(
        title=title,
        defaultextension=default_ext,
        filetypes=filetypes,
    )
    return path if path else None


def _write_text(path: str, content: str) -> None:
    with open(path, "w", encoding="utf-8", newline="") as f:
        f.write(content)


def _write_json(path: str, data: dict) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _write_csv(path: str, headers: list, rows: list) -> None:
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerows(rows)


def _write_docx(path: str, title: str, paragraphs: list) -> None:
    """写入简单 Word 报告。

    paragraphs 为字符串列表，空字符串表示新段落（留空一行）。
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise ImportError("请安装 python-docx: pip install python-docx") from e

    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    for para in paragraphs:
        if para == "":
            doc.add_paragraph()
        elif para.startswith("# "):
            doc.add_heading(para[2:], level=2)
        elif para.startswith("## "):
            doc.add_heading(para[3:], level=3)
        else:
            doc.add_paragraph(para)

    doc.save(path)


# --------------------------------------------------------------------------- #
# 基础分析导出
# --------------------------------------------------------------------------- #


def export_basic_result(res: BasicResult, parent=None) -> None:
    """导出基础分析结果。"""
    path = _ask_path(
        "导出基础分析结果",
        ".txt",
        [
            ("文本文件", "*.txt"),
            ("CSV 表格", "*.csv"),
            ("JSON 数据", "*.json"),
            ("Word 文档", "*.docx"),
        ],
    )
    if not path:
        return

    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".csv":
            _export_basic_csv(path, res)
        elif ext == ".json":
            _export_basic_json(path, res)
        elif ext == ".docx":
            _export_basic_docx(path, res)
        else:
            _export_basic_txt(path, res)
        messagebox.showinfo("导出成功", f"已保存：{path}")
    except Exception as e:
        messagebox.showerror("导出失败", str(e))


def _export_basic_txt(path: str, res: BasicResult) -> None:
    lines = [
        "汉英 NLP 分析工具 — 基础分析报告",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        res.summary(),
        "",
        "--- 分词结果 ---",
        "",
    ]
    for t in res.tokens:
        lines.append(f"{t.text}\t{t.pos}\t{t.lemma}")
    lines.extend(["", "--- 词频 Top 30 ---", ""])
    for w, c in res.freq.most_common(30):
        lines.append(f"{w}\t{c}")
    _write_text(path, "\n".join(lines))


def _export_basic_csv(path: str, res: BasicResult) -> None:
    rows = []
    for t in res.tokens:
        rows.append([t.text, t.pos, t.lemma, "是" if t.is_stop else "否"])
    _write_csv(path, ["词元", "词性", "词形还原", "停用词"], rows)


def _export_basic_json(path: str, res: BasicResult) -> None:
    data = {
        "summary": res.summary(),
        "tokens": [t.as_dict() for t in res.tokens],
        "frequency": dict(res.freq.most_common(100)),
        "pos_distribution": dict(res.pos_dist),
    }
    _write_json(path, data)


def _export_basic_docx(path: str, res: BasicResult) -> None:
    paragraphs = [
        "# 基础分析摘要",
        res.summary(),
        "",
        "## 分词结果",
    ]
    for t in res.tokens[:200]:
        paragraphs.append(f"{t.text}  ({t.pos}, {t.lemma})")
    if len(res.tokens) > 200:
        paragraphs.append(f"（仅显示前 200 条，共 {len(res.tokens)} 条）")
    paragraphs.extend(["", "## 词频 Top 30"])
    for w, c in res.freq.most_common(30):
        paragraphs.append(f"{w}: {c}")
    _write_docx(path, "基础分析报告", paragraphs)


# --------------------------------------------------------------------------- #
# 句法/语义导出
# --------------------------------------------------------------------------- #


def export_syntax_result(res: SyntaxResult, parent=None) -> None:
    """导出句法/语义分析结果。"""
    path = _ask_path(
        "导出句法/语义分析结果",
        ".txt",
        [
            ("文本文件", "*.txt"),
            ("CSV 表格", "*.csv"),
            ("JSON 数据", "*.json"),
            ("Word 文档", "*.docx"),
        ],
    )
    if not path:
        return

    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".csv":
            _export_syntax_csv(path, res)
        elif ext == ".json":
            _export_syntax_json(path, res)
        elif ext == ".docx":
            _export_syntax_docx(path, res)
        else:
            _export_syntax_txt(path, res)
        messagebox.showinfo("导出成功", f"已保存：{path}")
    except Exception as e:
        messagebox.showerror("导出失败", str(e))


def _export_syntax_txt(path: str, res: SyntaxResult) -> None:
    lines = [
        "汉英 NLP 分析工具 — 句法/语义分析报告",
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "--- 命名实体 ---",
    ]
    for e in res.ner:
        lines.append(f"{e['text']}\t{e['label']}")
    lines.extend(["", "--- 关键词 ---", ""])
    for w, weight in res.keywords:
        lines.append(f"{w}\t{weight:.4f}")
    lines.extend(["", "--- 依存关系 ---", ""])
    for d in res.dependencies:
        lines.append(
            f"{d['text']}\t{d['pos']}\t{d['dep']}\t→\t{d['head_text']}"
        )
    lines.extend(["", "--- 情感分析 ---", ""])
    lines.append(f"情感：{res.sentiment['label']}")
    lines.append(f"得分：{res.sentiment['score']}")
    lines.append(f"原始值：{res.sentiment['raw']}")
    _write_text(path, "\n".join(lines))


def _export_syntax_csv(path: str, res: SyntaxResult) -> None:
    rows = [[d["text"], d["pos"], d["dep"], d["head_text"], d["head_pos"]]
            for d in res.dependencies]
    _write_csv(path, ["词元", "词性", "依存", " head", "head 词性"], rows)


def _export_syntax_json(path: str, res: SyntaxResult) -> None:
    data = {
        "ner": res.ner,
        "keywords": res.keywords,
        "dependencies": res.dependencies,
        "sentiment": res.sentiment,
        "pos_tags": res.pos_tags,
    }
    _write_json(path, data)


def _export_syntax_docx(path: str, res: SyntaxResult) -> None:
    paragraphs = [
        "# 句法/语义分析报告",
        "",
        "## 命名实体",
    ]
    for e in res.ner:
        paragraphs.append(f"{e['text']}  ({e['label']})")
    paragraphs.extend(["", "## 关键词"])
    for w, weight in res.keywords:
        paragraphs.append(f"{w}: {weight:.4f}")
    paragraphs.extend(["", "## 依存关系"])
    for d in res.dependencies:
        paragraphs.append(
            f"{d['text']} ({d['pos']}) —{d['dep']}→ {d['head_text']} ({d['head_pos']})"
        )
    paragraphs.extend(["", "## 情感分析"])
    paragraphs.append(f"情感：{res.sentiment['label']}")
    paragraphs.append(f"得分：{res.sentiment['score']}")
    _write_docx(path, "句法/语义分析报告", paragraphs)


# --------------------------------------------------------------------------- #
# 可读性/对齐导出
# --------------------------------------------------------------------------- #


def export_readability_result(res, parent=None) -> None:
    """导出可读性结果。"""
    path = _ask_path("导出可读性结果", ".txt",
                     [("文本文件", "*.txt"), ("Word 文档", "*.docx")])
    if not path:
        return
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            _write_docx(path, "可读性分析报告", ["# 可读性分析", res.summary()])
        else:
            _write_text(path, f"可读性分析\n\n{res.summary()}")
        messagebox.showinfo("导出成功", f"已保存：{path}")
    except Exception as e:
        messagebox.showerror("导出失败", str(e))


def export_alignment_result(res: Alignment, parent=None) -> None:
    """导出中英对齐结果。"""
    path = _ask_path("导出对齐结果", ".txt",
                     [("文本文件", "*.txt"), ("CSV 表格", "*.csv"), ("Word 文档", "*.docx")])
    if not path:
        return
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".csv":
            rows = [[i + 1, z, e] for i, (z, e) in enumerate(res.pairs)]
            _write_csv(path, ["序号", "中文", "英文"], rows)
        elif ext == ".docx":
            paragraphs = ["# 中英句子对齐结果", res.summary(), ""]
            for i, (z, e) in enumerate(res.pairs, 1):
                paragraphs.append(f"## 句对 {i}")
                paragraphs.append(f"中文：{z}")
                paragraphs.append(f"英文：{e}")
                paragraphs.append("")
            _write_docx(path, "中英对齐报告", paragraphs)
        else:
            lines = ["中英句子对齐结果", res.summary(), ""]
            for i, (z, e) in enumerate(res.pairs, 1):
                lines.append(f"[{i}] {z}")
                lines.append(f"    {e}")
                lines.append("")
            _write_text(path, "\n".join(lines))
        messagebox.showinfo("导出成功", f"已保存：{path}")
    except Exception as e:
        messagebox.showerror("导出失败", str(e))


# --------------------------------------------------------------------------- #
# 语言指纹导出
# --------------------------------------------------------------------------- #


def export_batch_result(results: list, parent=None) -> None:
    """导出批量分析结果。"""
    path = _ask_path(
        "导出批量分析结果",
        ".csv",
        [("CSV 表格", "*.csv"), ("JSON 数据", "*.json"), ("Word 文档", "*.docx")],
    )
    if not path:
        return
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".json":
            data = [
                {
                    "filename": r.filename,
                    "status": r.status,
                    "error": r.error,
                    "lang": r.lang,
                    "char_count": r.char_count,
                    "word_count": r.word_count,
                    "sentence_count": r.sentence_count,
                    "unique_words": r.unique_words,
                    "top_words": r.top_words,
                }
                for r in results
            ]
            _write_json(path, {"batch_results": data})
        elif ext == ".docx":
            paragraphs = ["# 批量文件分析报告", ""]
            ok = sum(1 for r in results if r.status == "ok")
            paragraphs.append(f"成功：{ok}/{len(results)}")
            paragraphs.append("")
            for r in results:
                if r.status == "ok":
                    top = ", ".join(f"{w}({c})" for w, c in r.top_words)
                    paragraphs.append(
                        f"{r.filename}: {r.lang}, {r.char_count} 字符, "
                        f"{r.word_count} 词元, {r.sentence_count} 句子, "
                        f"不重复词 {r.unique_words}; Top: {top}"
                    )
                else:
                    paragraphs.append(f"{r.filename}: 失败 — {r.error}")
            _write_docx(path, "批量文件分析报告", paragraphs)
        else:
            rows = []
            for r in results:
                if r.status == "ok":
                    top = ", ".join(f"{w}({c})" for w, c in r.top_words)
                    rows.append(
                        [r.filename, r.lang, r.char_count, r.word_count,
                         r.sentence_count, r.unique_words, top]
                    )
                else:
                    rows.append([r.filename, "ERROR", "", "", "", "", r.error])
            _write_csv(path, ["文件名", "语言", "字符数", "词元数", "句子数", "不重复词", "Top 词"], rows)
        messagebox.showinfo("导出成功", f"已保存：{path}")
    except Exception as e:
        messagebox.showerror("导出失败", str(e))


def export_kwic_result(lines: list, parent=None) -> None:
    """导出 KWIC 结果。"""
    path = _ask_path("导出 KWIC 结果", ".txt",
                     [("文本文件", "*.txt"), ("CSV 表格", "*.csv")])
    if not path:
        return
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".csv":
            rows = [[line.left, line.keyword, line.right, line.position] for line in lines]
            _write_csv(path, ["左侧上下文", "关键词", "右侧上下文", "位置"], rows)
        else:
            if not lines:
                _write_text(path, "未找到匹配结果。")
            else:
                max_left = max(len(line.left) for line in lines)
                out = [f"共 {len(lines)} 条匹配", ""]
                for line in lines:
                    out.append(f"{line.left.rjust(max_left)}  [{line.keyword}]  {line.right}")
                _write_text(path, "\n".join(out))
        messagebox.showinfo("导出成功", f"已保存：{path}")
    except Exception as e:
        messagebox.showerror("导出失败", str(e))


def export_fingerprint_result(res: FingerprintResult, parent=None) -> None:
    """导出语言指纹分析结果。"""
    path = _ask_path(
        "导出语言指纹报告",
        ".txt",
        [("文本文件", "*.txt"), ("Word 文档", "*.docx")],
    )
    if not path:
        return
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            paragraphs = [
                "# 语言指纹分析报告",
                res.summary(),
                "",
                res.verdict_detail(),
            ]
            _write_docx(path, "语言指纹分析报告", paragraphs)
        else:
            _write_text(path, res.summary() + "\n\n" + res.verdict_detail())
        messagebox.showinfo("导出成功", f"已保存：{path}")
    except Exception as e:
        messagebox.showerror("导出失败", str(e))
