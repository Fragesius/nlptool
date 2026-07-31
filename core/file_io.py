"""文件读取工具 — 支持多种格式的文本提取。

支持格式：
  .txt     纯文本（自动检测编码）
  .docx    Microsoft Word
  .pdf     PDF 文档
  .html    网页（去标签）
  .md      Markdown
  .rtf     Rich Text Format（基础）
  .json    JSON（读取为文本）
  .csv     CSV（读取为文本）
  .xml     XML（读取为文本）
"""

from __future__ import annotations

import re
import os
from html.parser import HTMLParser
from typing import List, Optional

from core.log import logger


# ── 编码检测（无 chardet 依赖）───────────────────────────
# 优先尝试常见中文编码，再回退到 utf-8/latin-1
_ENCODINGS = ["utf-8", "gbk", "gb2312", "gb18030", "utf-16", "big5", "latin-1"]


def _detect_encoding(filepath: str) -> str:
    """按优先级尝试各编码，找到第一个成功解码的。"""
    raw = open(filepath, "rb").read()
    for enc in _ENCODINGS:
        try:
            raw.decode(enc)
            return enc
        except (UnicodeDecodeError, LookupError):
            continue
    return "latin-1"  # 最差情况，不会抛异常


# ── TXT ─────────────────────────────────────────────────


def _read_txt(filepath: str) -> str:
    enc = _detect_encoding(filepath)
    with open(filepath, encoding=enc, errors="replace") as f:
        return f.read()


# ── DOCX ─────────────────────────────────────────────────


def _read_docx(filepath: str) -> str:
    """读取 Word 文档的全部文本，段落之间用换行分隔。"""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # 也读表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


# ── PDF ─────────────────────────────────────────────────


def _read_pdf(filepath: str) -> str:
    """读取 PDF 的全部文本。"""
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except ImportError:
        raise ImportError("请安装 PyPDF2: pip install PyPDF2")

    reader = PdfReader(filepath)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


# ── HTML ─────────────────────────────────────────────────


class _HTMLStripper(HTMLParser):
    """HTML 标签剥离器。"""

    def __init__(self):
        super().__init__()
        self.parts: List[str] = []
        self._skip = False

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style", "noscript", "meta", "link"):
            self._skip = True

    def handle_endtag(self, tag):
        if tag in ("script", "style", "noscript"):
            self._skip = False
        # 块级标签后加换行
        if tag in ("p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6",
                    "li", "tr", "section", "article", "header", "footer"):
            self.parts.append("\n")

    def handle_data(self, data):
        if not self._skip:
            text = data.strip()
            if text:
                self.parts.append(text + " ")

    def handle_entityref(self, name):
        # 常见 HTML 实体
        entity_map = {
            "nbsp": " ", "amp": "&", "lt": "<", "gt": ">",
            "quot": '"', "apos": "'", "mdash": "—", "ndash": "–",
            "ldquo": "\u201c", "rdquo": "\u201d", "lsquo": "\u2018",
            "rsquo": "\u2019", "hellip": "…",
        }
        self.parts.append(entity_map.get(name, ""))


def _read_html(filepath: str) -> str:
    """从 HTML 文件中提取纯文本。"""
    enc = _detect_encoding(filepath)
    with open(filepath, encoding=enc, errors="replace") as f:
        html = f.read()

    stripper = _HTMLStripper()
    stripper.feed(html)
    stripper.close()

    # 清理多余空白
    text = "".join(stripper.parts)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


# ── RTF ─────────────────────────────────────────────────


def _read_rtf(filepath: str) -> str:
    """基础 RTF 文本提取（无第三方库依赖）。

    剥离常见 RTF 控制字和组标记，保留可见文本。
    """
    with open(filepath, encoding="latin-1", errors="replace") as f:
        rtf = f.read()

    # 剥离 RTF 控制字（\word 形式）
    text = re.sub(r"\\[a-zA-Z]+\d*", "", rtf)
    # 剥离 \{ \} 组标记
    text = re.sub(r"[\\{}]", "", text)
    # 剥离 hex 编码 \'xx
    text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
    # 清理多余空白
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


# ── 主入口 ──────────────────────────────────────────────


# 扩展名 → 读取函数
_READERS = {
    ".txt": _read_txt,
    ".docx": _read_docx,
    ".pdf": _read_pdf,
    ".html": _read_html,
    ".htm": _read_html,
    ".md": _read_txt,
    ".rtf": _read_rtf,
    ".json": _read_txt,
    ".csv": _read_txt,
    ".xml": _read_txt,
    ".text": _read_txt,
}

# 文件对话框的过滤器
FILETYPES = [
    ("所有支持的格式", "*.txt;*.docx;*.pdf;*.html;*.htm;*.md;*.rtf;*.json;*.csv;*.xml"),
    ("文本文件", "*.txt"),
    ("Word 文档", "*.docx"),
    ("PDF 文档", "*.pdf"),
    ("网页", "*.html;*.htm"),
    ("Markdown", "*.md"),
    ("RTF 文档", "*.rtf"),
    ("所有文件", "*.*"),
]

# 支持的扩展名列表（用于提示）
SUPPORTED_EXTENSIONS = list(_READERS.keys())


# 文件大小限制（10 MB），超过则提示分段处理
_MAX_FILE_SIZE = 10 * 1024 * 1024


def read_file(filepath: str) -> str:
    """读取文件内容，返回纯文本。

    根据扩展名自动选择合适的读取器。
    支持的格式：txt, docx, pdf, html, htm, md, rtf, json, csv, xml

    Args:
        filepath: 文件路径。

    Returns:
        提取的纯文本。

    Raises:
        ValueError: 不支持的格式。
        FileNotFoundError: 文件不存在。
        OSError: 文件过大。
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"文件不存在: {filepath}")

    size = os.path.getsize(filepath)
    if size > _MAX_FILE_SIZE:
        size_mb = size / (1024 * 1024)
        raise OSError(
            f"文件过大（{size_mb:.1f} MB）。"
            f"当前限制为 {_MAX_FILE_SIZE // (1024 * 1024)} MB，"
            f"建议拆分文件或提取其中部分章节后再试。"
        )

    ext = os.path.splitext(filepath)[1].lower()
    reader = _READERS.get(ext)
    if reader is None:
        raise ValueError(
            f"不支持的格式: {ext}\n"
            f"支持的格式: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    try:
        return reader(filepath)
    except ImportError as e:
        logger.error("读取文件 %s 失败（缺少依赖）: %s", filepath, e)
        raise ImportError(f"缺少依赖库：{e}")
    except Exception as e:
        logger.error("读取文件 %s 失败: %s", filepath, e, exc_info=True)
        raise




def _strip_chars(text: str) -> int:
    """统计不含任何空白字符的字符数。"""
    return len(re.sub(r"\s", "", text))
def read_file_with_label(filepath: str) -> tuple:
    """读取文件，返回 (文本内容, 状态标签)。

    状态标签示例: "已加载 26,389 字符 (gbk)"

    Returns:
        (text, label) 元组。
    """
    text = read_file(filepath)
    if not text.strip():
        return ("", "⚠ 文件为空或未能提取文本")

    char_count = _strip_chars(text)
    filename = os.path.basename(filepath)
    return (text, f"✅ {filename}  —  {char_count:,} 字符")


def get_file_summary(filepath: str) -> str:
    """快速获取文件摘要（不读取全部内容）。

    Returns:
        格式: "文件名  —  大小"
    """
    if not os.path.exists(filepath):
        return "文件不存在"
    name = os.path.basename(filepath)
    size = os.path.getsize(filepath)
    if size < 1024:
        size_str = f"{size} B"
    elif size < 1024 * 1024:
        size_str = f"{size / 1024:.1f} KB"
    else:
        size_str = f"{size / (1024 * 1024):.1f} MB"
    return f"{name}  —  {size_str}"
